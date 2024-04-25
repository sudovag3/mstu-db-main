import os
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
import time
import datetime
from getpass import getpass
from mysql.connector import connect, Error
import re
from docx import Document

from config.md_start import *
from config.md_routes import *

from dotenv import load_dotenv
dotenv_path = os.path.join(os.path.dirname(__file__), '.env')
if os.path.exists(dotenv_path):
    load_dotenv(dotenv_path)

class Programm:
    def __init__(self,host,user,password,port):
        self.host=host
        self.user=user
        self.password=password
        self.port=port
        
        self.create_update=create_update
        self.create_form=create_form
        self.change_accept=change_accept
        self.create_table=create_table
        self.draw_table=draw_table
        self.send_query=send_query
        self.main_menu=main_menu
        self.create_project=create_project
        
        
        self.win=Tk()
        self.width=520
        self.height=720
        self.xs=7 # ширина символа
        self.ys=20 # высота символа
        
       # self.xs=20
     #   self.ys=50
        
        self.canvas=Canvas(self.win,width=self.width,height=self.height,bg='black')
        self.canvas.pack(side=LEFT,expand=True,fill=BOTH)
        self.win.title('Title')

        self.win.option_add("*tearOff", FALSE)

        start(self) # инициализация всех переменных (модуль md_start)
        
# обработка движения мыши и нажатий ЛКМ
    def motion(self,evt):
        self.x=evt.x
        self.y=evt.y
    def bpress1(self,evt):
        x=evt.x
        y=evt.y
        for tr in self.d_zone:
            if self.x>=tr[0] and self.x<=tr[2] and self.y>=tr[1] and self.y<=tr[3]:
                
                self.tr=int(tr[4])
                if self.bd==4:
                    self.route_change('project')
                else:
                    self.route_change('update')
                
# валидаторы даты и bool                
    def date_is_valid(self,newval):
        #result=  re.match("^\+{0,1}\d{0,11}$",
       # result=  re.match("^\d{0,4}\-{0,1}\d{0,2}\-{0,1}\d{0,2}\d{0,2}\:{0,1}\d{0,2}\:{0,1}\d{0,2}$", newval) is not None
       result=  re.match("^\d{0,4}\-{0,1}\d{0,2}\-{0,1}\d{0,2}$", newval) is not None
       if not result and len(newval) <= 10 and self.route!='update':
            self.mprint("Дата должна быть записана в формате yyyy-mm-dd")
       return result
    def bool_is_valid(self,newval):
       result=  re.match("^\d{0,1}$", newval) is not None
       if not result and len(newval) <= 1:
            self.mprint("Переменная Boolean должна быть записана как 0 или 1")
       return result
       
# обработчики кнопок форм
    def accept_update(self):
        #
        tb=[self.titles[self.bd]]
        vst1=-1
        vst2=-1
        vst3=-1
        vst4=-1
        vst5=-1

        if 'Категория' in tb[0]:
            vst1=tb[0].index('Категория')-1
            query='SELECT * FROM '+self.tbs[0]
            data2=self.my_sql(query,True)
            dct1={}
            for d in data2:
                dct1[str(d[1])]=d[0]
        if 'Оборудование' in tb[0]:
            vst2=tb[0].index('Оборудование')-1
            query='SELECT * FROM '+self.tbs[3]
            data2=self.my_sql(query,True)
            dct2={}
            for d in data2:
                dct2[str(d[1])]=d[0]
        if 'Руководитель' in tb[0]:
            vst3=tb[0].index('Руководитель')-1
            query='SELECT * FROM '+self.tbs[2]
            data2=self.my_sql(query,True)
            dct3={}
            for d in data2:
                dct3[str(d[2])]=d[0]
        if 'Проект' in tb[0]:
            vst4=tb[0].index('Проект')
            query='SELECT * FROM '+self.tbs[4]
            data2=self.my_sql(query,True)
            dct4={}
            for d in data2:
                dct4[str(d[0])]=str(d[0])
        if 'Заказчик' in tb[0]:
            vst5=tb[0].index('Заказчик')
            query='SELECT * FROM '+self.tbs[5]
            data2=self.my_sql(query,True)
            dct5={}
            for d in data2:
                dct5[str(d[1])]=d[0]
        #
        arr=[]
        for i in range(len(self.obj)):
            a=self.obj[i].get()
            #
            if i==vst1:
                arr.append(dct1.get(a,a))
            elif i==vst2:
                arr.append(dct2.get(a,a))
            elif i==vst3:
                arr.append(dct3.get(a,a))
            elif i==vst4:
                arr.append(dct4.get(a,a))
            elif i==vst5:
                arr.append(dct5.get(a,a))
            else:
                arr.append(a)
            #
        try:
            arr=self.change_accept(self,arr)
            for a in arr:
                if a=='': a='null'
            st1=self.titles2[self.bd][1:-1].split(',')
            st2=arr.split(',')
            st3=''
            for i in range(len(st1)):
                st3+=st1[i]+'='+st2[i]+', '
            st3=st3[:-2]
            query='UPDATE '+self.tbs[self.bd]+' SET '
            query+=st3
            query+=' WHERE '+self.ids[self.bd]+'='+str(self.tr)
            self.my_sql(query,True)
            self.route_change('view')
        except Error as e:
            self.mprint(e)
            
    def accept_delete(self):
        st1=self.titles2[self.bd][1:-1].split(',')
        query='DELETE FROM '+self.tbs[self.bd]
        query+=' WHERE '+self.ids[self.bd]++'='+str(self.tr)
        self.my_sql(query,True)
        self.route_change('view')

    def accept_back(self):
        self.route_change('view')
       
    def button_update(self):
        self.route_change('update')
        
    def accept_generate(self):
        name='Отчеты/Отчет_'+str(time.time()//(60*60)%10000)+'_'+str(self.generate[0][1][0])+'.docx'
        titles=['Проект','Договора','Сотрудники','Оборудование']
        document = Document()
        for i in range(len(self.generate)):
            tb=self.generate[i]
            title=titles[i]
            document.add_heading(title)
            table = document.add_table(rows=len(tb), cols=len(tb[0]))
            for i in range(len(tb)):
                row = table.rows[i]
                for j in range(len(tb[i])):
                    row.cells[j].text = tb[i][j]
        
        document.save(name)

        
        self.mprint('Отчет сохранен в папку')
        
        
    def accept_create(self):
        #
        tb=[self.titles[self.bd]]
        vst1=-1
        vst2=-1
        vst3=-1
        vst4=-1
        vst5=-1

        if 'Категория' in tb[0]:
            vst1=tb[0].index('Категория')-1
            query='SELECT * FROM '+self.tbs[0]
            data2=self.my_sql(query,True)
            dct1={}
            for d in data2:
                dct1[str(d[1])]=d[0]
        if 'Оборудование' in tb[0]:
            vst2=tb[0].index('Оборудование')-1
            query='SELECT * FROM '+self.tbs[3]
            data2=self.my_sql(query,True)
            dct2={}
            for d in data2:
                dct2[str(d[1])]=d[0]
        if 'Руководитель' in tb[0]:
            vst3=tb[0].index('Руководитель')-1
            query='SELECT * FROM '+self.tbs[2]
            data2=self.my_sql(query,True)
            dct3={}
            for d in data2:
                dct3[str(d[2])]=d[0]
        if 'Проект' in tb[0]:
            vst4=tb[0].index('Проект')
            query='SELECT * FROM '+self.tbs[4]
            data2=self.my_sql(query,True)
            dct4={}
            for d in data2:
                dct4[str(d[0])]=str(d[0])
        if 'Заказчик' in tb[0]:
            vst5=tb[0].index('Заказчик')
            query='SELECT * FROM '+self.tbs[5]
            data2=self.my_sql(query,True)
            dct5={}
            for d in data2:
                dct5[str(d[1])]=d[0]
        #
        arr=[]
        for i in range(len(self.obj)):
            a=self.obj[i].get()
            self.obj[i].delete(0,END)
            #
            if i==vst1:
                arr.append(dct1.get(a,a))
            elif i==vst2:
                arr.append(dct2.get(a,a))
            elif i==vst3:
                arr.append(dct3.get(a,a))
            elif i==vst4:
                arr.append(dct4.get(a,a))
            elif i==vst5:
                arr.append(dct5.get(a,a))
            else:
                arr.append(a)
            #
        try:
            arr=self.change_accept(self,arr)
            query='INSERT INTO '+self.tbs[self.bd]+' '
            query+=self.titles2[self.bd]+'\n VALUES \n'
            query+='('+arr+')'
            self.my_sql(query,True)
        except:
            self.mprint('Ошибка типа переменных')
    # смена режима
    def route_change(self,name):
        self.obj=[]
        for t in self.tableim2:
            self.canvas.delete(t)
        self.tableim2=[]
        self.d_zone=[]
        
        self.route=name
        
        if name=='view':
            tb=self.create_table(self)
            self.draw_table(self,10,10,tb)
        elif name=='create':
            self.create_form(self,self.titles[self.bd].copy())
        elif name=='update':
            self.create_update(self,self.titles[self.bd].copy())
        elif name=='query':
            self.send_query(self)
        elif name=='':
            self.main_menu(self)
        elif name=='project':
            self.create_project(self)
      # сообшения об ошибках
    def mprint(self,txt):
        messagebox.showinfo("GUI Python", txt)
   # работа с MySQL
    def start_sql(self):
    	rz=0
    	try:
    		self.connection=connect(
        	host=self.host,
        	user=self.user,
        	password=self.password,
            port=self.port
    		)
    		rz=1
    	except Error as e:
    		self.mprint(e)
    	return rz
    	
    def create_db(self,name):
        try:
        	query = "CREATE DATABASE "+name
        	with self.connection.cursor() as cursor:
        		cursor.execute(query)
        except Error as e:
        	#self.mprint(e)
        	pass
        self.db=name
        try:
        	self.connection=connect(
        	host=self.host,
        	user=self.user,
        	password=self.password,
        	database=self.db,
            port=self.port
    		)
        except Error as e:
            self.mprint(e)
    def create_tables(self):
        tbls=[
        '''
        create table Category (
	category_id int not null primary key AUTO_INCREMENT COMMENT 'Код категории', 
    category_name varchar(30) COMMENT 'Название категории'  , 
    equipment_code int COMMENT 'Код оборудования'  )
        ''',
        '''
        create table Department (
	department_id int not null primary key AUTO_INCREMENT COMMENT 'Код департамента',
    department_name varchar(50) COMMENT 'Название департамента'  ,
    category_id int COMMENT 'Код категории'  ,
    FOREIGN KEY (category_id) REFERENCES Category(category_id)
)
        ''',
        '''
        create table Employee (
	employee_id int not null primary key AUTO_INCREMENT COMMENT 'Код сотрудника'  ,
    category_id int not null COMMENT 'Код категории'  ,
    employee_full_name varchar(50) not null COMMENT 'ФИО сотрудника'  ,
    hire_date datetime not null COMMENT 'Дата приёма'  ,
    fire_date datetime COMMENT 'Дата увольнения'  ,
    copyright_certificates_number int unsigned COMMENT 'Число авторских свидетельств'  ,
    outstaff_company_name varchar(50) COMMENT 'Наименование аутстафф-компании'  ,
    head_position_possible bool not null COMMENT 'Возможность занимания руководящей должности' 
)
        ''',
        '''
        create table Equipment(
	equipment_id int not null primary key AUTO_INCREMENT COMMENT 'Код оборудования',
    equipment_name varchar(50) not null COMMENT 'Название оборудования',
    category_id int not null COMMENT 'Код категории',
    FOREIGN KEY (category_id) REFERENCES Category(category_id)
)
        ''',
        '''
        create table Project (
	project_id int not null primary key AUTO_INCREMENT COMMENT 'Код проекта',
    project_materials varchar(300) COMMENT 'Материалы проекта',
    equipment_id int COMMENT 'Код оборудования',
    head_employee_id int not null COMMENT 'Код руководителя',
    category_id int not null COMMENT 'Код категории',
    
    FOREIGN KEY (category_id) REFERENCES Category(category_id),
    FOREIGN KEY (equipment_id) REFERENCES Equipment(equipment_id),
    FOREIGN KEY (head_employee_id) REFERENCES Employee(employee_id)
    
)
        ''',
        '''
        create table Customer (
	customer_id int not null primary key AUTO_INCREMENT COMMENT 'Код заказчика',
    customer_full_name varchar(50) not null COMMENT 'ФИО заказчика',
    customer_passport_data varchar(150) not null COMMENT 'Паспортные данные заказчика',
    customer_address varchar(50) COMMENT 'Адрес заказчика',
    customer_phone varchar(20) COMMENT 'Номер телефона заказчика'
)
        ''',
        '''
        create table Contract(
	contract_id int not null primary key COMMENT 'Код договора',
    project_id int not null COMMENT 'Код проекта',
    customer_id int not null COMMENT 'Код заказчика',
    deal_cost int unsigned not null COMMENT 'Стоимость работ',
    work_finish_date datetime not null COMMENT 'Дата окончания работ',
    
    FOREIGN KEY (project_id) REFERENCES Project(project_id),
    FOREIGN KEY (customer_id) REFERENCES Customer(customer_id)
)
        '''
        ]
        for tbl in tbls:
            self.my_sql(tbl,False)
    def my_sql(self,query,esend):
        try:
            with self.connection.cursor() as cursor:
            	cursor.execute(query)
            	result = cursor.fetchall()
            	return result
            		
        except Error as e:
        	if esend:
        	    self.mprint(e)
        	    

# блок команд для верхнего меню
    def send_query1(self):
        self.que=0
        self.route_change('query')
    def send_query2(self):
        self.que=1
        self.route_change('query')    
    def send_query3(self):
        self.que=2
        self.route_change('query')
    def send_query4(self):
        self.que=3
        self.route_change('query')
    def send_query5(self):
        self.que=4
        self.route_change('query')
    def send_query6(self):
        self.que=5
        self.route_change('query')        
    def send_query7(self):
        self.que=6
        self.route_change('query')        
    def send_query8(self):
        self.que=7
        self.route_change('query')                        
    def send_query9(self):
        self.que=8
        self.route_change('query')                                        
    def send_query10(self):
        self.que=9
        self.route_change('query')
    def send_query11(self):
        self.que=10
        self.route_change('query')       
    def send_query12(self):
        self.que=11
        self.route_change('query')                       
    def send_query13(self):
        self.que=12
        self.route_change('query') 
        
    def bd1_view(self):
        self.bd=0
        self.route_change('view')

    def bd2_view(self):
        self.bd=1
        self.route_change('view')

    def bd3_view(self):
        self.bd=2
        self.route_change('view')
      

    def bd4_view(self):
        self.bd=3
        self.route_change('view')
        

    def bd5_view(self):
        self.bd=4
        self.route_change('view')
       

    def bd6_view(self):
        self.bd=5
        self.route_change('view')
       

    def bd7_view(self):
        self.bd=6
        self.route_change('view')
     

    def bd1_create(self):
        self.bd=0
        self.route_change('create')
       

    def bd2_create(self):
        self.bd=1
        self.route_change('create')
      

    def bd3_create(self):
        self.bd=2
        self.route_change('create')
        

    def bd4_create(self):
        self.bd=3
        self.route_change('create')
        

    def bd5_create(self):
        self.bd=4
        self.route_change('create')
        

    def bd6_create(self):
        self.bd=5
        self.route_change('create')

    def bd7_create(self):
        self.bd=6
        self.route_change('create')
        
    def end_command(self):
    	self.end=1

    def button_mm(self):
        self.route_change('')
        
# описание верхнего меню
    def create_menu(self):
        main_menu = Menu()
         
        
        
        main_menu.add_command(label='В главное меню',command=self.button_mm)

        titles=[
        'Проекты',
        'Названия категорий',
        'На руководящую должность',
        'Договора в работе',
        'Количество сотрудников',
        'Кол-во активных проектов от конкретного заказчика',
        'Нагрузка на команды',
        'Кол-во проектов с несколькими командами',
        'Названия команд, работающих над несколькими проектами',
        'ФИО сотрудника и название команды, работающих над несколькими проектами',
        'Кол-во сотрудников, которые не пришли из аутстафф-компаний',
        'Общее кол-во денег, которое принёс каждый заказчик',
        'Кол-во заказов от каждого заказчика'
        ]

        menu1 = Menu()
        menu1.add_command(label=titles[0], command=self.send_query1)
        menu1.add_command(label=titles[1], command=self.send_query2)
        menu1.add_command(label=titles[2], command=self.send_query3)
        menu1.add_command(label=titles[3], command=self.send_query4)
        menu1.add_command(label=titles[4], command=self.send_query5)
        menu1.add_command(label=titles[5], command=self.send_query6)
        menu1.add_command(label=titles[6], command=self.send_query7)
        menu1.add_command(label=titles[7], command=self.send_query8)
        menu1.add_command(label=titles[8], command=self.send_query9)
        menu1.add_command(label=titles[9], command=self.send_query10)
        menu1.add_command(label=titles[10], command=self.send_query11)
        menu1.add_command(label=titles[11], command=self.send_query12)
        menu1.add_command(label=titles[12], command=self.send_query13)
        main_menu.add_cascade(label="Запросы", menu=menu1)
        
        main_menu.add_command(label='Выход',command=self.end_command)
         
        self.win.config(menu=main_menu)
# описание полос прокрутки
    def create_scroll(self):
        
        self.canvas.config(scrollregion=(0,0,520,720))
        hbar=Scrollbar(self.canvas,orient=HORIZONTAL)
        hbar.pack(side=BOTTOM,fill=X)
        hbar.config(command=self.canvas.xview)
        vbar=Scrollbar(self.canvas,orient=VERTICAL)
        vbar.pack(side=RIGHT,fill=Y)
        vbar.config(command=self.canvas.yview)
        self.canvas.config(xscrollcommand=hbar.set,yscrollcommand=vbar.set)
        self.win.geometry(str(self.width)+"x"+str(self.height))
# изменение области прокрутки
    def set_scroll(self,x,y):
        self.canvas.config(scrollregion=(0,0,x,y))

    def mainloop(self):
        while self.end==0:
  
            self.overdraw()
            self.win.update()
        self.win.destroy()
        
 
        
    def overdraw(self):
        for t in self.tableim:
            self.canvas.delete(t)
        self.tableim=[]
        
        for tr in self.d_zone:
            if self.x>=tr[0] and self.x<=tr[2] and self.y>=tr[1] and self.y<=tr[3]:
                self.tableim.append(self.canvas.create_rectangle(tr[0],tr[1],tr[2],tr[3],fill='grey'))
                self.canvas.tag_lower(self.tableim[-1])


host=os.environ.get('DB_HOST')
user=os.environ.get('DB_USER')
password=os.environ.get('DB_PASSWORD') or ''
port=os.environ.get('DB_PORT')

programm=Programm(host,user,password, port)
programm.create_menu()
programm.create_scroll()
programm.route_change('')



programm.start_sql()
programm.create_db('rgr_db')
programm.create_tables()

programm.mainloop()
